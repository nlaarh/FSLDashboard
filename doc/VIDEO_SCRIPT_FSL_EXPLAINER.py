"""
Generate the FSL System Explainer Word Document
Run: python3 doc/VIDEO_SCRIPT_FSL_EXPLAINER.py
Output: doc/FSL_System_Explainer_Script.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_title(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    return h

def add_scene(number, title, duration):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'SCENE {number}: {title}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'⏱ Duration: ~{duration}')
    run2.italic = True
    run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    return p

def add_visual(text):
    p = doc.add_paragraph()
    run = p.add_run(f'🎬 VISUAL: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
    run2 = p.add_run(text)
    run2.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
    return p

def add_narration(text):
    p = doc.add_paragraph()
    run = p.add_run(f'🎙 NARRATION: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
    run2 = p.add_run(text)
    run2.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f'📝 NOTE: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
    run2 = p.add_run(text)
    run2.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    run2.italic = True
    return p

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')
run = p.add_run('AAA Western & Central New York')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Field Service Lightning (FSL)\nHow the System Works')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nExecutive Video Script & System Guide')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\nFor narration via ElevenLabs AI Voice')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
run = p.add_run('\nEstimated video length: 12-15 minutes')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
run = p.add_run('\nMarch 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════

add_title('Table of Contents', level=1)
toc = [
    'Scene 1: The Big Picture — AAA\'s Service Territory (45 sec)',
    'Scene 2: Meet the Players — Garages, Drivers, and Territories (90 sec)',
    'Scene 3: Fleet vs Contractor — Two Different Worlds (60 sec)',
    'Scene 4: A Call Comes In — Following a Real Request (90 sec)',
    'Scene 5: The System Picks a Driver — How Assignment Works (120 sec)',
    'Scene 6: The Home Address Problem — Why the Scheduler Struggles (90 sec)',
    'Scene 7: When Things Go Wrong — Cascading and Declining (60 sec)',
    'Scene 8: The 15-Minute Rule — Balancing Speed vs Cost (60 sec)',
    'Scene 9: GPS — The Eyes of the System (60 sec)',
    'Scene 10: Scoring — How We Measure Performance (90 sec)',
    'Scene 11: What We\'re Building — FleetPulse AI Dashboard (60 sec)',
    'Scene 12: The Path Forward (45 sec)',
    '',
    'Appendix A: Complete Data Model & Field Reference',
    'Appendix B: Scoring Dimensions & Weights',
    'Appendix C: Work Types & Skill Hierarchy',
    'Appendix D: System Users & Integration Points',
]
for item in toc:
    if item:
        doc.add_paragraph(item, style='List Number' if 'Scene' in item else 'List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SCENE 1: THE BIG PICTURE
# ══════════════════════════════════════════════════════════════════════════════

add_scene(1, 'THE BIG PICTURE — AAA\'s Service Territory', '45 seconds')

add_visual('Animated map of New York State zooming into Western & Central NY. Territory boundaries light up one by one across the region — Buffalo, Rochester, Syracuse, Utica. Each territory glows with a colored border. Numbers pop up: "74 active territories", "~800 calls/day", "71 fleet drivers + 300 contractor drivers".')

add_narration(
    'AAA Western and Central New York covers a massive service area — from the shores of Lake Erie in Buffalo, '
    'across the Finger Lakes, through Syracuse, and all the way to the Adirondacks. '
    'Every day, roughly 800 people call for roadside help: dead batteries, flat tires, lockouts, tows. '
    'Each call needs to reach the right garage, the right driver, at the right time. '
    'The system that makes this happen is called Salesforce Field Service Lightning — or FSL. '
    'Let\'s follow a single call through the system to understand how it works, and where it struggles.'
)

add_note('Animation should show the full WCNY region with all 74 active territory polygons. Use the KML boundary data from ServiceTerritory records. Highlight the scale — this is NOT a single city, it spans hundreds of miles.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SCENE 2: MEET THE PLAYERS
# ══════════════════════════════════════════════════════════════════════════════

add_scene(2, 'MEET THE PLAYERS — Garages, Drivers, and Territories', '90 seconds')

add_visual('Map zooms into the Buffalo area. Three layers appear one at a time: (1) Territory polygons with names like "076DO - Transit Auto Detail", "727 - M&R Automotive". (2) Garage pin icons at each territory center. (3) Driver dots — blue for fleet, pink for contractor — scattered across the territories.')

add_narration(
    'The system is organized into three layers. '
    'First: Territories. A territory is a geographic zone — think of it as a coverage area on the map. '
    'Each territory has boundaries defined by GPS coordinates, like zip code regions. '
    'There are 74 active territories across our region. \n\n'
    'Second: Garages. Each territory is served by a garage — a physical shop where drivers are based. '
    'Some garages are AAA fleet shops staffed by our own employees. '
    'Others are independent contractors — towing companies like Transit Auto Detail or Ricci and Sons — '
    'that we dispatch through a separate system called Towbook. \n\n'
    'Third: Drivers. This is where it gets interesting. '
    'Fleet drivers are AAA employees. They use the Salesforce Field Service mobile app on their phones, '
    'which sends their GPS location to the system every few minutes. '
    'The scheduler can see where they are in real time. \n\n'
    'Contractor drivers are different. They work for the towing company, not for AAA. '
    'They don\'t use the FSL app. They don\'t send GPS. '
    'When we dispatch a call to a contractor garage, the garage\'s own dispatcher decides which of their drivers to send. '
    'We only find out the driver arrived when Towbook sends us a status update.'
)

add_visual('Split screen animation: LEFT shows a fleet driver with a phone icon (FSL app) sending GPS pings to a satellite/cloud. RIGHT shows a contractor driver with no phone icon — a question mark where GPS should be.')

add_narration(
    'This creates a fundamental split in our data. '
    'For fleet calls, we know everything: where the driver is, how fast they\'re moving, when they arrive. '
    'For contractor calls, we\'re partially blind. We know the call went to the garage, '
    'and we know when the driver eventually shows up — but we don\'t know who was sent or where they started from.'
)

add_note('Show the Priority Matrix concept: each territory has a PRIMARY garage (rank 1) and BACKUP garages (rank 2, 3...). Animate arrows from a territory polygon to its primary garage, then dotted arrows to backup garages. Use the ERS_Territory_Priority_Matrix__c data.')

p = doc.add_paragraph()
p.add_run('Territory → Garage Assignment (Priority Matrix)').bold = True
doc.add_paragraph(
    'Every territory has a ranked list of garages that can serve it, stored in the '
    'ERS_Territory_Priority_Matrix__c object:\n\n'
    '  • Rank 1 = Primary garage (gets the call first)\n'
    '  • Rank 2 = First backup (if primary declines or is full)\n'
    '  • Rank 3+ = Additional cascades\n\n'
    'The priority can also be filtered by work type. For example, a territory might use Garage A '
    'for tows (they have flatbeds) but Garage B for battery calls (they\'re closer for light service).\n\n'
    'Key fields:\n'
    '  • ERS_Parent_Service_Territory__c → the zone/territory\n'
    '  • ERS_Spotted_Territory__c → the garage assigned\n'
    '  • ERS_Priority__c → rank number (1 = primary)\n'
    '  • ERS_Worktype__c → optional work type filter'
)

p = doc.add_paragraph()
p.add_run('Driver → Territory Assignment').bold = True
doc.add_paragraph(
    'Drivers are linked to territories through ServiceTerritoryMember records:\n\n'
    '  • ServiceResourceId → the driver\n'
    '  • ServiceTerritoryId → the territory/garage\n'
    '  • TerritoryType: P (Primary), S (Secondary), R (Relocation)\n'
    '  • Address fields (Street, City, State, PostalCode, Latitude, Longitude)\n\n'
    'CRITICAL ISSUE: Of 501 ServiceTerritoryMember records, ZERO have a street address populated. '
    'Only 65 have geocoded latitude/longitude. This is the root cause of the scheduler\'s biggest problem, '
    'which we\'ll explain in Scene 6.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SCENE 3: FLEET VS CONTRACTOR
# ══════════════════════════════════════════════════════════════════════════════

add_scene(3, 'FLEET VS CONTRACTOR — Two Different Worlds', '60 seconds')

add_visual('Two-panel comparison. LEFT: "AAA Fleet" — blue brand color. Shows a driver icon with GPS signal, FSL app on phone, AAA-branded truck. RIGHT: "Contractor (Towbook)" — pink color. Shows a tow truck icon with NO GPS signal, a separate Towbook system icon, an independent company logo.')

add_narration(
    'Let\'s be clear about the two types of garages, because they operate completely differently. \n\n'
    'Fleet garages are AAA-owned. We employ the drivers, we own the trucks, and we control the technology. '
    'Fleet drivers run the FSL mobile app, which gives us real-time GPS. '
    'The FSL Enhanced Scheduler can automatically assign calls to fleet drivers based on their location, '
    'skills, and availability. When a fleet driver arrives on scene, the timestamp in Salesforce is accurate. \n\n'
    'Contractor garages are independent towing companies. They have their own drivers, their own trucks, '
    'their own internal dispatch. When AAA sends a call to a contractor, it goes through Towbook — '
    'a separate dispatch platform. The contractor\'s dispatcher decides who to send. '
    'We have no GPS visibility into their drivers. And here\'s a critical gotcha: '
    'the arrival time that Salesforce records for contractor calls is NOT the real arrival time. '
    'Towbook writes a fake estimated time into the ActualStartTime field when the call is completed. '
    'To get the real arrival, we have to dig into the ServiceAppointmentHistory and find when the status '
    'changed to "On Location" — that timestamp is the actual moment the driver showed up.'
)

p = doc.add_paragraph()
p.add_run('Side-by-Side Comparison Table:').bold = True

table = doc.add_table(rows=8, cols=3)
table.style = 'Medium Grid 1 Accent 1'
headers = ['Aspect', 'Fleet (Field Services)', 'Contractor (Towbook)']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['Dispatch Method', 'ERS_Dispatch_Method__c = "Field Services"', 'ERS_Dispatch_Method__c = "Towbook"'],
    ['Who assigns driver', 'FSL Scheduler + AAA dispatchers', 'Towbook facility dispatcher (external)'],
    ['GPS tracking', 'Real-time via FSL mobile app', 'NONE — no GPS data sent to Salesforce'],
    ['Real arrival time', 'ActualStartTime is accurate', 'FAKE — must use SAHistory "On Location" event'],
    ['Driver visibility', 'Full: name, location, skills, availability', 'Partial: truck ID, estimated times only'],
    ['Drivers in Salesforce', 'Real count from ServiceTerritoryMember', '1 placeholder per garage (e.g. "Towbook-076DO")'],
    ['Cost control', 'Full — AAA payroll, fuel, vehicles', 'Per-call contract rate — variable cost'],
]
for r, row in enumerate(data):
    for c, val in enumerate(row):
        table.rows[r + 1].cells[c].text = val

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SCENE 4: A CALL COMES IN
# ══════════════════════════════════════════════════════════════════════════════

add_scene(4, 'A CALL COMES IN — Following a Real Request', '90 seconds')

add_visual('Animation: A car icon appears on the map near Exit 42 on I-90 (between Syracuse and Utica). A red "HELP!" pulse radiates from the car. A phone icon appears — the member is calling AAA. The call routes to a call center building icon.')

add_narration(
    'It\'s 2:30 PM on a Tuesday. Maria is driving on the I-90 Thruway near Exit 42, between Syracuse and Utica. '
    'Her car won\'t start after a rest stop. She calls AAA. \n\n'
    'The call center agent creates a Service Appointment in Salesforce. '
    'This is the moment the clock starts. The CreatedDate timestamp — 2:30 PM — will be used to measure '
    'every performance metric: response time, SLA compliance, PTA accuracy. Everything traces back to this moment.'
)

add_visual('The ServiceAppointment record appears on screen, filling in field by field with a typewriter animation:\n'
    '  • Status: Dispatched\n'
    '  • Created: 2:30 PM\n'
    '  • Location: Exit 42, I-90 (43.05°N, 75.38°W)\n'
    '  • Work Type: Battery\n'
    '  • Territory: assigned to Syracuse zone\n'
    '  • PTA: 45 minutes (promised to member)')

add_narration(
    'The system captures everything: where Maria is, what type of help she needs — in this case, a dead battery — '
    'and which territory zone she\'s in. \n\n'
    'Then, the Mulesoft integration checks the PTA settings for this territory and work type. '
    'PTA stands for Promised Time of Arrival. For battery calls in this Syracuse zone, the setting is 45 minutes. '
    'This is what the call center tells Maria: "We\'ll have someone there within 45 minutes." \n\n'
    'That 45-minute promise is now recorded on the Service Appointment. '
    'If we deliver in 35 minutes — great, we beat the promise. '
    'If it takes 60 minutes — we broke our word to the member. '
    'The PTA accuracy score tracks how often we keep that promise.'
)

add_note('Show the PTA lookup: Territory + WorkType → PTA Minutes. Animate the Mulesoft integration reading the ERS_Service_Appointment_PTA__c record. Show the 45-minute timer starting on screen, counting down.')

p = doc.add_paragraph()
p.add_run('What the system records at call creation:').bold = True
doc.add_paragraph(
    '  • ServiceAppointment.CreatedDate = 2:30:00 PM (THE clock start)\n'
    '  • ServiceAppointment.Status = "Dispatched"\n'
    '  • ServiceAppointment.Latitude / Longitude = member GPS\n'
    '  • ServiceAppointment.ServiceTerritoryId = zone lookup\n'
    '  • ServiceAppointment.WorkType = "Battery"\n'
    '  • ServiceAppointment.ERS_PTA__c = 45 (promised minutes)\n'
    '  • ServiceAppointment.ERS_Dispatch_Method__c = assigned later (Fleet or Towbook)\n'
    '  • ServiceAppointment.ERS_Parent_Territory__c = parent zone for cascade'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SCENE 5: THE SYSTEM PICKS A DRIVER
# ══════════════════════════════════════════════════════════════════════════════

add_scene(5, 'THE SYSTEM PICKS A DRIVER — How Assignment Works', '120 seconds')

add_visual('Map view centered on Maria\'s location (Exit 42). The territory boundary highlights. The PRIMARY garage for this zone lights up (Rank 1 from Priority Matrix). An arrow connects the territory to the garage. Then, driver dots appear inside the garage\'s territory — blue dots for fleet drivers with GPS, grey dots for drivers without GPS.')

add_narration(
    'Now the system needs to find a driver. Here\'s the decision tree. \n\n'
    'Step one: Territory to Garage. The Priority Matrix tells the system which garage is primary for this zone. '
    'Let\'s say it\'s Garage 414 — a fleet garage in Syracuse. This garage has 5 drivers assigned to it. \n\n'
    'Step two: The FSL Enhanced Scheduler evaluates those drivers. '
    'It checks four things.'
)

add_visual('Four boxes appear one by one, each with an icon:\n'
    '  1. 🔧 SKILLS — Can this driver handle a battery call?\n'
    '  2. 📍 LOCATION — How far is this driver from Maria?\n'
    '  3. ⏰ AVAILABILITY — Is this driver free, or already on another call?\n'
    '  4. 🚫 ABSENCES — Is this driver marked unavailable (break, PTO, no GPS)?')

add_narration(
    'First, Skills. Not every driver can handle every call type. '
    'A battery specialist can only do batteries and jumpstarts. '
    'A light service driver can handle batteries, tires, lockouts, and fuel. '
    'A tow driver can do everything — tows plus all light service plus batteries. '
    'The system checks the driver\'s ServiceResourceSkill records against the WorkType\'s SkillRequirement. '
    'If there\'s no match, that driver is excluded. \n\n'
    'Second, Location — and this is where the biggest problem lives. '
    'The scheduler calculates travel time from the driver\'s "home base" address. '
    'This address comes from the ServiceTerritoryMember record — the link between the driver and the garage. '
    'We\'ll come back to why this is a disaster in the next scene. \n\n'
    'Third, Availability. The system checks if the driver is already assigned to another call. '
    'If they\'re currently on a tow that won\'t finish for another hour, they\'re not a candidate. \n\n'
    'Fourth, Absences. Resource Absence records mark drivers as unavailable. '
    'This includes breaks, PTO, call-outs, but also a critical workaround we use: '
    'marking drivers who don\'t have GPS as "absent" so the scheduler doesn\'t try to assign them blind.'
)

add_visual('Animation shows the 5 drivers being evaluated:\n'
    '  • Driver 1 (Mike): Skills ✓, Location ✓, Free ✓, No absence ✓ → CANDIDATE\n'
    '  • Driver 2 (Lisa): Skills ✓, Location ✓, On another call ✗ → EXCLUDED\n'
    '  • Driver 3 (Tom): Skills ✗ (tow-only, no battery) → EXCLUDED\n'
    '  • Driver 4 (Jay): Skills ✓, Location ✓, Free ✓, Marked absent (no GPS) ✗ → EXCLUDED\n'
    '  • Driver 5 (Sara): Skills ✓, Location ✓, Free ✓, No absence ✓ → CANDIDATE\n'
    '\n'
    'Mike and Sara remain. System picks Mike based on "closest to home base".')

add_narration(
    'Out of five drivers, only two pass all four checks. '
    'The scheduler picks Mike — because his home base address calculates as closer to Maria. '
    'The system creates an AssignedResource record linking Mike to Maria\'s Service Appointment. '
    'The SchedStartTime is set — this records when the assignment was made. \n\n'
    'Now here\'s what the dashboard measures from this point: '
    'Dispatch Speed is the time from call creation to assignment. '
    'If Maria called at 2:30 and Mike was assigned at 2:35, that\'s a 5-minute dispatch speed. '
    'If it took 20 minutes to find a driver, that\'s 20 minutes of the member\'s 45-minute promise already burned.'
)

add_note('Show a timeline bar at the bottom: 2:30 (call created) → 2:35 (assigned) → ?? (driver arrives). The 45-min SLA timer continues counting. Show "5 min dispatch speed" label on the assignment gap.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SCENE 6: THE HOME ADDRESS PROBLEM
# ══════════════════════════════════════════════════════════════════════════════

add_scene(6, 'THE HOME ADDRESS PROBLEM — Why the Scheduler Struggles', '90 seconds')

add_visual('Split screen. LEFT: What the scheduler THINKS — a driver icon sitting at a house icon labeled "Home: Syracuse shop, 12 miles from Maria". RIGHT: What\'s ACTUALLY happening — the same driver icon is on the I-90, just 2 miles from Maria, having just finished another call. A big red X over the scheduler\'s view. A green checkmark over reality.')

add_narration(
    'This is the single biggest challenge with the FSL Enhanced Scheduler, and it\'s critical to understand. \n\n'
    'AAA roadside drivers are NOT like Amazon delivery drivers. They don\'t start from home each morning '
    'and return at night. They\'re constantly on the road — finishing one call, driving to the next. '
    'At any given moment, a driver could be anywhere in their territory. \n\n'
    'But the scheduler doesn\'t know that. It calculates travel time from the driver\'s "home base" — '
    'the address stored in the ServiceTerritoryMember record. And here\'s the devastating fact: '
    'out of 501 driver-territory records, zero — literally zero — have a street address populated. '
    'Only 65 out of 501 even have a geocoded latitude and longitude.'
)

add_visual('A database table animates on screen, showing ServiceTerritoryMember records:\n'
    '  | Driver      | Territory        | Street | City | Lat/Lon    |\n'
    '  | Mike        | Syracuse         | NULL   | NULL | NULL       |\n'
    '  | Lisa        | Syracuse         | NULL   | NULL | 43.0, -76.1|\n'
    '  | Tom         | Rochester        | NULL   | NULL | NULL       |\n'
    '  | ...         | ...              | NULL   | NULL | NULL       |\n'
    '\n'
    'Red highlights on all the NULL cells. A counter shows: "0 of 501 addresses populated".')

add_narration(
    'The scheduler is essentially dispatching blind. It has work rules like '
    '"Maximum Travel From Home is 30 minutes" — but if there\'s no home address, '
    'it can\'t calculate travel at all. The driver either gets excluded entirely, '
    'or assigned based on incomplete data. \n\n'
    'Meanwhile, 23 of our 71 fleet drivers DO have live GPS reporting through the FSL mobile app. '
    'The scheduler has access to this field — LastKnownLatitude, LastKnownLongitude — '
    'but it doesn\'t use it for assignment. It always defaults to the home base address. '
    'This is a Salesforce platform limitation, not a configuration choice. \n\n'
    'The result? The scheduler might assign a driver 30 miles away '
    'when there\'s another driver sitting 2 miles from the member. '
    'The member waits longer. The truck burns more fuel. '
    'And the 45-minute SLA gets blown.'
)

add_visual('Map animation showing the failure case:\n'
    '  1. Maria\'s car (red pin) at Exit 42\n'
    '  2. Mike\'s REAL location (blue dot) — 2 miles away on I-90\n'
    '  3. Mike\'s HOME BASE (grey house icon) — 30 miles away in Watertown\n'
    '  4. Scheduler draws a line from house → Maria: "35 min travel"\n'
    '  5. Reality line from Mike → Maria: "4 min travel"\n'
    '  6. Big text: "31 MINUTES WASTED"')

add_narration(
    'Before we implemented the Resource Absence workaround, auto-assignment was at zero percent. '
    'Nobody was getting assigned by the system — dispatchers had to do everything manually. '
    'After we started using Resource Absences to filter out drivers without GPS, '
    'auto-assignment jumped to 83 percent. But it\'s a workaround, not a fix. '
    'The fundamental problem remains: the scheduler doesn\'t know where drivers actually are.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SCENE 7: WHEN THINGS GO WRONG
# ══════════════════════════════════════════════════════════════════════════════

add_scene(7, 'WHEN THINGS GO WRONG — Cascading and Declining', '60 seconds')

add_visual('Map showing Maria\'s call. The primary garage (Rank 1) flashes red — "DECLINED". An arrow redirects to the backup garage (Rank 2). The timer at the bottom keeps counting — 45 min SLA burning down.')

add_narration(
    'Not every assignment goes smoothly. Sometimes a garage declines the call. '
    'Maybe all their drivers are busy. Maybe the call type doesn\'t match their equipment. '
    'Maybe they just can\'t get there in time. \n\n'
    'When this happens, the system cascades. It looks up the Priority Matrix and tries the next ranked garage. '
    'Rank 1 declined? Try Rank 2. Rank 2 declined? Try Rank 3. '
    'Each cascade adds time — and Maria is still waiting on the Thruway.'
)

add_visual('The cascade flowchart animates:\n'
    '  Maria\'s Call → Garage 414 (Rank 1, Fleet) → DECLINED (all drivers busy)\n'
    '                → Garage 076DO (Rank 2, Contractor) → ACCEPTED\n'
    '                   → Towbook dispatches their driver\n'
    '                   → 10 more minutes pass before someone is en route\n'
    '\n'
    'Timer shows: 2:30 (call) → 2:35 (first attempt) → 2:40 (decline) → 2:45 (cascade to 076DO) → 2:52 (driver en route)')

add_narration(
    'In our system, we track cascading through the "spotting number." '
    'Spotting number 1 means the call was handled on the first try. '
    'Spotting number 2 means one redirect. Spotting 3 or higher means the call bounced '
    'between multiple garages — a cascade. \n\n'
    'Our data shows about 26 percent of calls hit spotting 3 or higher — meaning over a quarter '
    'of calls are being redirected multiple times before someone picks them up. '
    'Each redirect burns precious minutes off that 45-minute promise.'
)

add_note('Decline reasons are stored in ERS_Facility_Decline_Reason__c. Show the top reasons as a bar chart animation. Spotting data comes from ERS_Spotting_Number__c on ServiceAppointment.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SCENE 8: THE 15-MINUTE RULE
# ══════════════════════════════════════════════════════════════════════════════

add_scene(8, 'THE 15-MINUTE RULE — Balancing Speed vs Cost', '60 seconds')

add_visual('Map showing Maria\'s location and two driver options:\n'
    '  • Driver A (blue): 5 miles away, ETA 12 min — the CLOSEST by distance\n'
    '  • Driver B (green): 15 miles away, ETA 8 min — the FASTEST (on highway, less traffic)\n'
    'A balance scale animation: LEFT side = "Cost (miles)" / RIGHT side = "Speed (ETA)"')

add_narration(
    'Here\'s the optimization challenge we\'re trying to solve. \n\n'
    'In a perfect world, you always send the closest driver by distance. '
    'Fewer miles means less fuel, less truck wear, lower cost. '
    'But the closest driver isn\'t always the fastest. '
    'A driver 5 miles away on back roads might take 20 minutes. '
    'A driver 15 miles away on the highway might arrive in 10. \n\n'
    'We call this the 15-Minute Rule. '
    'If the closest driver\'s estimated arrival is within 15 minutes of the fastest driver — '
    'send the closest. The cost savings are worth the slightly longer wait. '
    '15 minutes is the tax we\'re willing to pay to optimize driving costs. \n\n'
    'But if the gap is more than 15 minutes? Send the faster driver. '
    'Because at that point, the member has been waiting too long, '
    'and the cost savings don\'t justify the service impact.'
)

add_visual('Two examples animate side by side:\n'
    '  EXAMPLE 1: Closest ETA = 25 min, Fastest ETA = 20 min\n'
    '    → Gap = 5 min (< 15 min threshold)\n'
    '    → SEND CLOSEST — save miles ✓\n\n'
    '  EXAMPLE 2: Closest ETA = 50 min, Fastest ETA = 30 min\n'
    '    → Gap = 20 min (> 15 min threshold)\n'
    '    → SEND FASTEST — member waiting too long ✓')

add_note('This rule is the core dispatch optimization philosophy. Currently the scheduler cannot implement this because it lacks real-time GPS-based ETAs. The FleetPulse dashboard surfaces this analysis to dispatchers manually.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SCENE 9: GPS
# ══════════════════════════════════════════════════════════════════════════════

add_scene(9, 'GPS — The Eyes of the System', '60 seconds')

add_visual('Dashboard view of the Fleet Drivers tile from FleetPulse. Numbers fill in: 71 Drivers, 32% GPS Tracking, 23 On GPS. The bar chart shows green/amber/red segments. Then zoom into the map showing blue dots (with GPS) and grey question marks (without).')

add_narration(
    'GPS is the foundation of everything. Without it, the scheduler is blind, '
    'the closest-driver analysis can\'t run, and dispatchers are guessing. \n\n'
    'Right now, we have 71 active fleet drivers. Of those, only 23 — about 32 percent — '
    'are reporting GPS in real time. 25 drivers have NEVER reported GPS. '
    'They\'re marked as active fleet drivers in Salesforce, but they may not have the FSL app installed, '
    'or it may not be running. \n\n'
    'For contractors, it\'s zero. None of the 300-plus contractor drivers report GPS to Salesforce. '
    'Their location is a black box until they arrive on scene. \n\n'
    'This is why we built FleetPulse — to visualize exactly this problem. '
    'The dashboard shows fleet driver positions in real time, tracks GPS freshness, '
    'and surfaces recommendations that the scheduler can\'t make on its own.'
)

add_note('GPS data source: ServiceResource.LastKnownLatitude / LastKnownLongitude / LastKnownLocationDate. Historical GPS trail stored in ServiceResourceHistory — 4.25 million records. This enables reconstructing "where was the driver when the call was assigned?" for forensic analysis.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SCENE 10: SCORING
# ══════════════════════════════════════════════════════════════════════════════

add_scene(10, 'SCORING — How We Measure Performance', '90 seconds')

add_visual('A garage scorecard fills in dimension by dimension, like a report card:\n'
    '  Garage: 414 — Syracuse Fleet\n'
    '  45-Min SLA: 78% (target: 100%) — AMBER\n'
    '  Completion Rate: 92% (target: 95%) — AMBER\n'
    '  Customer Satisfaction: 85% (target: 82%) — GREEN\n'
    '  Median Response: 52 min (target: 45 min) — RED\n'
    '  PTA Accuracy: 71% (target: 90%) — RED\n'
    '  Could Not Wait: 4% (target: 3%) — AMBER\n'
    '  Dispatch Speed: 8 min (target: 5 min) — AMBER\n'
    '  Decline Rate: 1% (target: 2%) — GREEN\n'
    '  ──────────────────────────────\n'
    '  COMPOSITE: 72 / 100 — Grade C')

add_narration(
    'Every garage is scored across 8 dimensions, rolling 28 days of data. '
    'This is the performance scorecard. \n\n'
    'The most important metric — weighted at 30 percent — is the 45-Minute SLA Hit Rate. '
    'This measures: of all completed calls, how many had a driver arrive within 45 minutes of the call? '
    'That\'s AAA\'s standard. If you called at 2:30, someone should be at your car by 3:15. \n\n'
    'Completion Rate, at 15 percent, tracks how many calls actually get completed versus canceled or no-showed. '
    'Customer Satisfaction, also 15 percent, comes from post-call surveys — '
    'specifically the "Totally Satisfied" percentage. \n\n'
    'Median Response Time looks at the actual minutes from call to arrival. '
    'PTA Accuracy checks: did we deliver within the time we promised? '
    'Could Not Wait Rate catches when members give up and leave before the driver arrives. '
    'Dispatch Speed measures how fast we assign a driver. '
    'And Decline Rate tracks how often a garage says "no" to a call.'
)

add_visual('A weight breakdown pie chart appears:\n'
    '  30% — 45-Min SLA\n'
    '  15% — Completion Rate\n'
    '  15% — Customer Satisfaction\n'
    '  10% — Median Response Time\n'
    '  10% — PTA Accuracy\n'
    '  10% — Could Not Wait\n'
    '   5% — Dispatch Speed\n'
    '   5% — Decline Rate')

add_narration(
    'Each dimension gets a score from 0 to 100, weighted, and combined into a composite grade. '
    'A and above is 90 — the garage is excellent. C is 70 — needs improvement. '
    'Below 60 is an F — serious performance problems that need immediate attention. \n\n'
    'The FleetPulse Insights panel uses these scores to generate specific, actionable recommendations: '
    'which garages need help, where to rebalance drivers, and where the PTA promise doesn\'t match reality.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SCENE 11: WHAT WE'RE BUILDING
# ══════════════════════════════════════════════════════════════════════════════

add_scene(11, 'WHAT WE\'RE BUILDING — FleetPulse AI Dashboard', '60 seconds')

add_visual('Screen recording / mockup of the FleetPulse dashboard:\n'
    '  1. Command Center — live map with territories, driver dots, open calls\n'
    '  2. Dispatch Split — System vs Dispatcher percentage\n'
    '  3. Closest Driver analysis — did we send the nearest one?\n'
    '  4. Capacity Alerts — garages overwhelmed right now\n'
    '  5. Insights tab — AI-generated recommendations\n'
    '  6. Garage Dashboard — scorecard for each garage\n'
    '  7. PTA Advisor — are our promises accurate?\n'
    '  8. Territory Matrix — rebalancing recommendations')

add_narration(
    'FleetPulse is the intelligence layer we\'re building on top of Salesforce. '
    'It doesn\'t replace the FSL scheduler — it fills the gaps the scheduler can\'t. \n\n'
    'The Command Center shows everything happening right now: '
    'which territories have open calls, where drivers are, who\'s over capacity. '
    'It compares System dispatch versus Dispatcher dispatch in real time — '
    'so we can measure how well the auto-scheduler is performing versus human decisions. \n\n'
    'The Closest Driver analysis looks at every completed call and asks: '
    'did we send the nearest available driver? If not, how many extra miles did we waste? '
    'This is the metric that directly measures dispatch quality. \n\n'
    'The Insights engine — powered by AI — analyzes 28 days of data across all garages '
    'and generates actionable recommendations: where to move drivers, '
    'which territories need rebalancing, and which garages need operational attention. \n\n'
    'And the AI chat assistant lets dispatchers and managers ask questions in plain English: '
    '"Which garage has the longest wait times right now?" '
    '"How is Transit Auto Detail performing this week?"'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SCENE 12: THE PATH FORWARD
# ══════════════════════════════════════════════════════════════════════════════

add_scene(12, 'THE PATH FORWARD', '45 seconds')

add_visual('A roadmap graphic with three phases:\n'
    '  PHASE 1 (NOW): FleetPulse Dashboard — visibility into what\'s happening\n'
    '  PHASE 2 (NEXT): GPS Coverage Push — get all fleet drivers on the FSL app\n'
    '  PHASE 3 (FUTURE): Smart Assignment — custom logic using real GPS for dispatch')

add_narration(
    'The path forward has three phases. \n\n'
    'Phase one is what we have now: FleetPulse gives us visibility. '
    'We can see the problems — where drivers are wasted, where cascading is high, '
    'where the scheduler is making bad decisions. '
    'Visibility is the first step to improvement. \n\n'
    'Phase two is fixing GPS coverage. If we can get all 71 fleet drivers reporting GPS consistently, '
    'the scheduler\'s decisions improve dramatically — even with the home base limitation. '
    'More GPS means more data means better recommendations from FleetPulse. \n\n'
    'Phase three is the endgame: custom assignment logic that uses real-time GPS, '
    'not home base addresses, to calculate true closest driver. '
    'We have 4.25 million historical GPS records in ServiceResourceHistory. '
    'The data exists. The question is building the logic to use it. \n\n'
    'Every minute we shave off average response time is a member who gets home sooner, '
    'a truck that burns less fuel, and a promise we keep.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX A: COMPLETE DATA MODEL
# ══════════════════════════════════════════════════════════════════════════════

add_title('Appendix A: Complete Data Model & Field Reference', level=1)

objects = [
    ('ServiceAppointment', 'The call record — one per member request', [
        ('Id', 'Unique Salesforce ID'),
        ('AppointmentNumber', 'Human-readable call number'),
        ('Status', 'Dispatched → Assigned → Accepted → En Route → On Location → In Progress → Completed / Canceled / No-Show / Unable to Complete'),
        ('CreatedDate', 'Call creation timestamp — THE clock start for all metrics'),
        ('ActualStartTime', 'Driver arrival (Fleet=accurate, Towbook=FAKE)'),
        ('SchedStartTime', 'When assignment was made'),
        ('ERS_PTA__c', 'Promised time of arrival in minutes'),
        ('ERS_Dispatch_Method__c', '"Field Services" (fleet) or "Towbook" (contractor)'),
        ('ServiceTerritoryId', 'Which territory/zone this call belongs to'),
        ('WorkType.Name', 'Tow, Battery, Tire, Lockout, Fuel, Winch Out, etc.'),
        ('Latitude / Longitude', 'Member location (where the car is)'),
        ('ERS_Parent_Territory__c', 'Parent zone for cascade/priority matrix lookup'),
        ('ERS_Spotting_Number__c', 'Cascade count: 1=first try, 3+=cascaded'),
        ('ERS_Cancellation_Reason__c', 'Why call was canceled'),
        ('ERS_Facility_Decline_Reason__c', 'Why garage declined the call'),
        ('Off_Platform_Driver__c', 'Towbook driver contact (contractors only)'),
        ('Off_Platform_Truck_Id__c', 'Towbook truck ID (contractors only)'),
    ]),
    ('ServiceAppointmentHistory', 'Audit trail — every field change on an SA', [
        ('ServiceAppointmentId', 'Links back to the SA'),
        ('Field', 'Which field changed (e.g. "Status")'),
        ('OldValue / NewValue', 'Previous and new values'),
        ('CreatedDate', 'When the change happened — used for Towbook real arrival'),
        ('CreatedById', 'Who made the change (system user or human)'),
    ]),
    ('ServiceTerritory', 'A territory/garage — geographic zone or physical shop', [
        ('Id / Name', 'ID and display name (e.g. "076DO - TRANSIT AUTO DETAIL")'),
        ('Latitude / Longitude', 'Center point of the territory'),
        ('IsActive', 'Whether currently operational'),
        ('ERS_Facility_Account__c', 'Linked Account record for phone, address, dispatch info'),
        ('ParentTerritoryId', 'Parent zone (for hierarchy)'),
    ]),
    ('ServiceTerritoryMember', 'Links drivers to territories — THE home address record', [
        ('ServiceResourceId', 'Which driver'),
        ('ServiceTerritoryId', 'Which territory'),
        ('TerritoryType', 'P=Primary, S=Secondary, R=Relocation'),
        ('Street / City / State / PostalCode', 'Home address — 0 of 501 populated'),
        ('Latitude / Longitude', 'Geocoded home — only 65 of 501 have data'),
        ('EffectiveStartDate / EffectiveEndDate', 'Active assignment period'),
    ]),
    ('ServiceResource', 'A driver or technician', [
        ('Name', 'Driver name'),
        ('ERS_Driver_Type__c', 'Fleet Driver / On-Platform Contractor / Off-Platform Contractor'),
        ('LastKnownLatitude / Longitude', 'Real-time GPS from FSL mobile app'),
        ('LastKnownLocationDate', 'When GPS last updated'),
        ('IsActive', 'Active in the system'),
        ('ResourceType', '"T" for Technician'),
    ]),
    ('AssignedResource', 'Links a driver to a specific call', [
        ('ServiceResourceId', 'Which driver was assigned'),
        ('ServiceAppointmentId', 'Which call'),
        ('CreatedBy.Name', 'Who made the assignment (system or dispatcher name)'),
    ]),
    ('WorkType', 'Types of service calls', [
        ('Name', 'Tow, Battery, Jumpstart, Tire, Lockout, Fuel Delivery, Winch Out, PVS, Flat Bed, Tow Drop-Off'),
    ]),
    ('ERS_Territory_Priority_Matrix__c', 'Territory → Garage routing (cascade order)', [
        ('ERS_Parent_Service_Territory__c', 'The zone requesting service'),
        ('ERS_Spotted_Territory__c', 'The garage that can serve it'),
        ('ERS_Priority__c', 'Rank: 1=primary, 2+=backup cascade'),
        ('ERS_Worktype__c', 'Optional: filter by work type'),
    ]),
    ('ERS_Service_Appointment_PTA__c', 'PTA settings — promised arrival times', [
        ('ERS_Service_Territory__c', 'Which garage/territory'),
        ('ERS_Type__c', 'Tow / Winch / Battery / Light Service'),
        ('ERS_Minutes__c', 'Promised arrival in minutes'),
    ]),
    ('ResourceAbsence', 'Marks drivers as unavailable', [
        ('ResourceId', 'Which driver'),
        ('Start / End', 'Unavailable period'),
        ('Type', 'Real-Time Location / Break / Call out / PTO/Vacation'),
    ]),
    ('ServiceResourceHistory', 'Historical GPS trail — 4.25M records', [
        ('ServiceResourceId', 'Which driver'),
        ('Field', '"LastKnownLatitude" or "LastKnownLongitude"'),
        ('NewValue', 'The GPS coordinate value'),
        ('CreatedDate', 'Timestamp of the GPS update'),
    ]),
]

for obj_name, obj_desc, fields in objects:
    p = doc.add_paragraph()
    run = p.add_run(f'{obj_name}')
    run.bold = True
    run.font.size = Pt(12)
    p2 = doc.add_paragraph(obj_desc)
    p2.runs[0].italic = True

    table = doc.add_table(rows=len(fields) + 1, cols=2)
    table.style = 'Light List Accent 1'
    table.rows[0].cells[0].text = 'Field'
    table.rows[0].cells[1].text = 'Description'
    for i, (fname, fdesc) in enumerate(fields):
        table.rows[i + 1].cells[0].text = fname
        table.rows[i + 1].cells[1].text = fdesc
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX B: SCORING
# ══════════════════════════════════════════════════════════════════════════════

add_title('Appendix B: Scoring Dimensions & Weights', level=1)

doc.add_paragraph('The Garage Performance Scorecard uses 8 dimensions, weighted and combined into a composite 0-100 score (rolling 28 days):')

table = doc.add_table(rows=9, cols=5)
table.style = 'Medium Grid 1 Accent 1'
headers = ['Dimension', 'Weight', 'Target', 'Direction', 'Formula']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

scoring_data = [
    ['45-Min SLA Hit Rate', '30%', '100%', 'Higher = better', 'count(ATA ≤ 45min) ÷ count(valid ATAs)'],
    ['Completion Rate', '15%', '95%', 'Higher = better', 'count(Completed) ÷ count(all SAs)'],
    ['Customer Satisfaction', '15%', '82%', 'Higher = better', 'count(Totally Satisfied) ÷ count(surveys)'],
    ['Median Response Time', '10%', '45 min', 'Lower = better', 'median(all ATA values)'],
    ['PTA Accuracy', '10%', '90%', 'Higher = better', 'count(ATA ≤ ERS_PTA__c) ÷ count(both defined)'],
    ['"Could Not Wait" Rate', '10%', '3%', 'Lower = better', 'count(CNW cancellations) ÷ count(all SAs)'],
    ['Dispatch Speed', '5%', '5 min', 'Lower = better', 'median(SchedStartTime - CreatedDate)'],
    ['Facility Decline Rate', '5%', '2%', 'Lower = better', 'count(declined) ÷ count(all SAs)'],
]
for r, row in enumerate(scoring_data):
    for c, val in enumerate(row):
        table.rows[r + 1].cells[c].text = val

doc.add_paragraph()
doc.add_paragraph('Grading: A ≥ 90, B ≥ 80, C ≥ 70, D ≥ 60, F < 60')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX C: WORK TYPES & SKILLS
# ══════════════════════════════════════════════════════════════════════════════

add_title('Appendix C: Work Types & Skill Hierarchy', level=1)

doc.add_paragraph('Drivers are classified by capability tier. Higher tiers can serve all lower-tier call types:')
doc.add_paragraph(
    '  TOW (most versatile)\n'
    '  ├── Can do: Tow, Flat Bed, Wheel Lift\n'
    '  ├── Plus all Light Service calls\n'
    '  └── Plus all Battery calls\n\n'
    '  LIGHT SERVICE\n'
    '  ├── Can do: Tire, Lockout, Locksmith, Winch Out, Fuel, PVS\n'
    '  └── Plus all Battery calls\n\n'
    '  BATTERY (least versatile)\n'
    '  └── Can do: Battery, Jumpstart only\n\n'
    '  TOW DROP-OFF (special)\n'
    '  └── Second leg of a tow pickup — excluded from all metrics'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX D: SYSTEM USERS
# ══════════════════════════════════════════════════════════════════════════════

add_title('Appendix D: System Users & Integration Points', level=1)

doc.add_paragraph('These system accounts appear in ServiceAppointmentHistory as the "who made the change" field:')

table = doc.add_table(rows=6, cols=3)
table.style = 'Light List Accent 1'
table.rows[0].cells[0].text = 'System User'
table.rows[0].cells[1].text = 'Role'
table.rows[0].cells[2].text = 'Classified As'

users = [
    ['IT System User', 'Runs flows, triggers, and the FSL Scheduler', 'System'],
    ['Automated Process', 'Salesforce Process Builder / Flow automation', 'System'],
    ['Mulesoft Integration', 'Creates SAs from external call sources, sets PTA', 'System'],
    ['Replicant Integration User', 'IVR/voice system (automated phone intake)', 'System'],
    ['Integrations Towbook', 'Towbook sends status updates + timestamps (no GPS)', 'Contractor'],
]
for r, row in enumerate(users):
    for c, val in enumerate(row):
        table.rows[r + 1].cells[c].text = val

doc.add_paragraph()
doc.add_paragraph(
    'When analyzing who dispatched a call:\n'
    '  • If AssignedResource.CreatedBy matches a system user → System dispatch\n'
    '  • If it\'s a human name → Dispatcher (manual) dispatch\n'
    '  • If ERS_Dispatch_Method__c = "Towbook" → Contractor dispatch'
)

# ── Save ──
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, 'FSL_System_Explainer_Script.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
